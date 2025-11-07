#!/usr/bin/env python3
"""
Generate End-of-Day (EOD) report as a Microsoft Word document.

This script generates professional EOD reports from structured YAML input,
properly replacing template content while preserving formatting.

Usage:
    # Generate from YAML file
    python generate_eod_report.py eod_notes.yaml

    # Dry-run mode (preview without saving)
    python generate_eod_report.py eod_notes.yaml --dry-run

    # Archive old reports
    python generate_eod_report.py --archive --days 30

Author: QA Team
"""

import sys
import argparse
import subprocess
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Any
import re

# Add parent to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from common.docx_utils import (
    get_template_path,
    get_report_output_path,
    get_report_archive_path,
    load_docx_safely,
    save_docx_safely,
    print_section_header
)

try:
    import yaml
except ImportError:
    print("ERROR: PyYAML is not installed.", file=sys.stderr)
    print("Please run: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed.", file=sys.stderr)
    print("Please run: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)


# Template configuration
TEMPLATE_NAME = "Reports end of the day highlights.docx"
SECTIONS = [
    "1. Product and Environment Tested",
    "2. Areas Covered During Testing",
    "3. New Bugs / QA Notes Identified",
    "4. Bug Fixes Verified",
    "5. Requirements / Stories Confirmed",
    "6. Pending / Next Steps",
    "7. Testing Status",
]


def get_user_fullname() -> str:
    """
    Get user's full name from git config or system.

    Returns:
        Full name of the user
    """
    try:
        result = subprocess.run(
            ['git', 'config', 'user.name'],
            capture_output=True,
            text=True,
            check=True
        )
        name = result.stdout.strip()
        if name:
            return name.replace(' ', '_')
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass

    # Fallback to environment or default
    import getpass
    return getpass.getuser()


def parse_date_flexible(date_str: str) -> datetime:
    """
    Parse date from multiple formats.

    Supported formats:
    - DD-MM-YYYY (06-11-2025)
    - YYYY-MM-DD (2025-11-06)
    - MM/DD/YYYY (11/06/2025)
    - Month DD, YYYY (November 6, 2025)
    - DD Month YYYY (6 November 2025)

    Args:
        date_str: Date string to parse

    Returns:
        datetime object

    Raises:
        ValueError: If date cannot be parsed
    """
    date_str = date_str.strip()

    formats = [
        "%d-%m-%Y",      # 06-11-2025
        "%Y-%m-%d",      # 2025-11-06
        "%m/%d/%Y",      # 11/06/2025
        "%B %d, %Y",     # November 6, 2025
        "%d %B %Y",      # 6 November 2025
        "%b %d, %Y",     # Nov 6, 2025
        "%d %b %Y",      # 6 Nov 2025
    ]

    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue

    raise ValueError(f"Unable to parse date: '{date_str}'. Supported formats: DD-MM-YYYY, YYYY-MM-DD, MM/DD/YYYY, 'Month DD, YYYY'")


def delete_paragraph(paragraph) -> None:
    """Delete a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)


def _is_section_heading(paragraph_text: str) -> bool:
    """Check if a paragraph is a section heading."""
    stripped = paragraph_text.strip()
    return bool(re.match(r'^\d+\.', stripped))


def _find_paragraph_index(doc: Document, text: str) -> int:
    """Find the index of a paragraph with exact text match."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    return -1


def _replace_content_in_section(
    doc: Document,
    section_heading: str,
    new_content_lines: List[str],
    use_bullets: bool = False
) -> None:
    """
    Replace content in a section.

    Args:
        doc: Document object
        section_heading: The exact heading text to find
        new_content_lines: List of content lines to insert
        use_bullets: If True, format lines as bullets
    """
    # Find section heading index
    start_idx = _find_paragraph_index(doc, section_heading)
    if start_idx < 0:
        raise ValueError(f"Section heading not found: '{section_heading}'")

    # Find where this section ends
    end_idx = len(doc.paragraphs)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if _is_section_heading(doc.paragraphs[i].text):
            end_idx = i
            break

    # Delete existing content
    paragraphs_to_delete = []
    for i in range(start_idx + 1, end_idx):
        paragraphs_to_delete.append(doc.paragraphs[i])

    for p in paragraphs_to_delete:
        delete_paragraph(p)

    # Insert new content
    heading_element = doc.paragraphs[start_idx]._element

    for line in new_content_lines:
        if use_bullets:
            try:
                new_p = doc.add_paragraph(style='List Bullet')
            except KeyError:
                new_p = doc.add_paragraph()
                new_p.add_run("• " + line)
                heading_element.addnext(new_p._element)
                heading_element = new_p._element
                continue
        else:
            new_p = doc.add_paragraph()

        new_p.add_run(line)
        heading_element.addnext(new_p._element)
        heading_element = new_p._element


def _replace_header_date(doc: Document, formatted_date: str) -> None:
    """Replace the date in the header while preserving formatting."""
    for p in doc.paragraphs:
        if "summarizing the testing activities and findings for" in p.text:
            for run in p.runs:
                if "<DATE>" in run.text:
                    run.text = run.text.replace("<DATE>", formatted_date)
                elif "summarizing the testing activities and findings for" in run.text:
                    p.text = f"Here is the end-of-day report summarizing the testing activities and findings for {formatted_date}."
                    break
            break


def validate_template_structure(doc: Document) -> None:
    """
    Validate that the template contains all required section headings.

    Raises:
        ValueError: If any sections are missing
    """
    missing_sections = []
    for section in SECTIONS:
        if _find_paragraph_index(doc, section) < 0:
            missing_sections.append(section)

    if missing_sections:
        raise ValueError(
            f"Template is missing required sections:\n" +
            "\n".join(f"  - {s}" for s in missing_sections)
        )


def load_yaml_input(yaml_path: Path) -> Dict[str, Any]:
    """
    Load and validate YAML input file.

    Args:
        yaml_path: Path to YAML input file

    Returns:
        Dictionary with parsed YAML data

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If YAML is invalid or missing required fields
    """
    if not yaml_path.exists():
        raise FileNotFoundError(f"Input file not found: {yaml_path}")

    try:
        with open(yaml_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
    except yaml.YAMLError as e:
        raise ValueError(f"Invalid YAML format: {e}")

    # Validate required fields
    required_fields = ['date', 'product', 'status']
    missing_fields = [f for f in required_fields if f not in data]
    if missing_fields:
        raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")

    return data


def format_product_section(data: Dict[str, Any]) -> List[str]:
    """Format product and environment section."""
    product = data.get('product', {})
    lines = [f"Product: {product.get('name', 'Hello Britannica')}"]

    if 'platforms' in product:
        lines.append("Environments:")
        lines.extend(product['platforms'])

    if 'roles_tested' in product:
        roles = ", ".join(product['roles_tested'])
        lines.append(f"Roles Tested: {roles}")

    return lines


def format_areas_section(data: Dict[str, Any]) -> List[str]:
    """Format areas covered section."""
    areas = data.get('areas_covered', [])
    if not areas:
        return ["Regression and exploratory testing were carried out on both platforms."]
    return areas


def format_bugs_section(data: Dict[str, Any]) -> List[str]:
    """Format new bugs section."""
    bugs = data.get('bugs', [])

    if not bugs:
        return ["No new bugs were reported today. Regression testing continued as planned."]

    lines = []
    for bug in bugs:
        if isinstance(bug, str):
            lines.append(bug)
        elif isinstance(bug, dict):
            severity = bug.get('severity', '')
            title = bug.get('title', '')
            severity_str = f"[{severity}] " if severity else ""
            lines.append(f"{severity_str}{title}")
            if 'description' in bug:
                lines.append(f"  Description: {bug['description']}")

    return lines


def format_bug_fixes_section(data: Dict[str, Any]) -> List[str]:
    """Format bug fixes verified section."""
    fixes = data.get('bug_fixes', [])

    if not fixes:
        return ["No previously reported bugs were verified or closed today."]

    lines = []
    for fix in fixes:
        if isinstance(fix, str):
            lines.append(fix)
        elif isinstance(fix, dict):
            bug_id = fix.get('bug_id', '')
            title = fix.get('title', '')
            status = fix.get('status', 'Verified')
            id_str = f"{bug_id}: " if bug_id else ""
            lines.append(f"{id_str}{title} - {status}")

    return lines


def format_requirements_section(data: Dict[str, Any]) -> List[str]:
    """Format requirements/stories confirmed section."""
    reqs = data.get('requirements', [])

    if not reqs:
        return ["No new requirements or user stories were confirmed today."]

    lines = []
    for req in reqs:
        if isinstance(req, str):
            lines.append(req)
        elif isinstance(req, dict):
            story_id = req.get('story_id', '')
            title = req.get('title', '')
            status = req.get('status', 'Confirmed')
            id_str = f"{story_id}: " if story_id else ""
            lines.append(f"{id_str}{title} - {status}")

    return lines


def format_next_steps_section(data: Dict[str, Any]) -> List[str]:
    """Format next steps section."""
    next_steps = data.get('next_steps', [])
    if not next_steps:
        return ["Continue with regression and exploratory testing in current focus areas."]
    return next_steps


def generate_eod_report(
    yaml_path: Path,
    output_path: Optional[Path] = None,
    dry_run: bool = False
) -> Optional[Path]:
    """
    Generate EOD report from YAML input.

    Args:
        yaml_path: Path to YAML input file
        output_path: Optional custom output path
        dry_run: If True, don't save the file

    Returns:
        Path to generated report (None if dry_run)
    """
    # Load and validate input
    print_section_header("Loading EOD Input Data")
    data = load_yaml_input(yaml_path)

    # Parse date
    date_obj = parse_date_flexible(data['date'])
    formatted_date = date_obj.strftime("%B %d, %Y")
    print(f"Report Date: {formatted_date}")

    # Get tester name
    tester_name = data.get('tester', {}).get('name', get_user_fullname())
    print(f"Tester: {tester_name}")

    # Load template
    print_section_header("Loading Template")
    template_path = get_template_path(TEMPLATE_NAME)
    doc = load_docx_safely(template_path)
    if doc is None:
        raise FileNotFoundError(f"Could not load template: {template_path}")

    print(f"Template: {template_path}")

    # Validate template structure
    validate_template_structure(doc)
    print("Template validation: OK")

    # Generate content
    print_section_header("Generating Report Content")

    # Update header date
    _replace_header_date(doc, formatted_date)

    # Replace each section
    _replace_content_in_section(doc, SECTIONS[0], format_product_section(data), use_bullets=False)
    _replace_content_in_section(doc, SECTIONS[1], format_areas_section(data), use_bullets=True)
    _replace_content_in_section(doc, SECTIONS[2], format_bugs_section(data), use_bullets=False)
    _replace_content_in_section(doc, SECTIONS[3], format_bug_fixes_section(data), use_bullets=False)
    _replace_content_in_section(doc, SECTIONS[4], format_requirements_section(data), use_bullets=False)
    _replace_content_in_section(doc, SECTIONS[5], format_next_steps_section(data), use_bullets=True)
    _replace_content_in_section(doc, SECTIONS[6], [data['status']], use_bullets=False)

    print("All sections populated successfully")

    # Determine output path
    if output_path is None:
        file_date = date_obj.strftime("%Y-%m-%d")
        filename = f"EOD_{file_date}_{tester_name.replace(' ', '_')}.docx"
        output_path = get_report_output_path(filename)

    # Save or preview
    if dry_run:
        print_section_header("Dry-Run Mode - Preview")
        print(f"Would save to: {output_path}")
        print(f"File size: {len(doc.element.xml)} bytes (estimated)")
        print("\nContent preview:")
        print(f"  Product: {data.get('product', {}).get('name', 'N/A')}")
        print(f"  Areas covered: {len(format_areas_section(data))} items")
        print(f"  Bugs reported: {len(data.get('bugs', []))} bugs")
        print(f"  Status: {data['status']}")
        return None
    else:
        print_section_header("Saving Report")
        if save_docx_safely(doc, output_path):
            print(f"SUCCESS: Report saved to: {output_path}")
            print(f"File size: {output_path.stat().st_size} bytes")
            return output_path
        else:
            raise IOError("Failed to save report")


def archive_old_reports(days_to_keep: int = 30, dry_run: bool = False) -> int:
    """
    Archive old EOD reports to archive directory.

    Args:
        days_to_keep: Number of days to keep reports in main folder
        dry_run: If True, don't actually move files

    Returns:
        Number of reports archived
    """
    from common.docx_utils import get_project_root

    reports_dir = get_project_root() / "documentation" / "reports"
    cutoff = datetime.now() - timedelta(days=days_to_keep)

    archived_count = 0

    print_section_header(f"Archiving Reports Older Than {days_to_keep} Days")
    print(f"Cutoff date: {cutoff.strftime('%Y-%m-%d')}")

    for eod_file in reports_dir.glob("EOD_*.docx"):
        # Get file modification time
        file_mtime = datetime.fromtimestamp(eod_file.stat().st_mtime)

        if file_mtime < cutoff:
            # Determine archive location (by month)
            year_month = file_mtime.strftime("%Y-%m")
            archive_path = get_report_archive_path(eod_file.name, year_month)

            if dry_run:
                print(f"[DRY-RUN] Would move: {eod_file.name} -> archive/{year_month}/")
            else:
                # Ensure archive directory exists
                archive_path.parent.mkdir(parents=True, exist_ok=True)
                # Move file
                eod_file.rename(archive_path)
                print(f"Archived: {eod_file.name} -> archive/{year_month}/")

            archived_count += 1

    if archived_count == 0:
        print("No reports to archive.")
    else:
        print(f"\nTotal reports archived: {archived_count}")

    return archived_count


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Generate EOD reports from structured YAML input",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate report from YAML
  python generate_eod_report.py eod_notes.yaml

  # Preview without saving
  python generate_eod_report.py eod_notes.yaml --dry-run

  # Archive old reports (older than 30 days)
  python generate_eod_report.py --archive --days 30

  # Preview archival
  python generate_eod_report.py --archive --days 30 --dry-run
        """
    )

    parser.add_argument(
        'input_file',
        nargs='?',
        type=Path,
        help='YAML input file with EOD data'
    )

    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Preview mode - don\'t save files'
    )

    parser.add_argument(
        '--archive',
        action='store_true',
        help='Archive old EOD reports'
    )

    parser.add_argument(
        '--days',
        type=int,
        default=30,
        help='Days to keep reports before archiving (default: 30)'
    )

    parser.add_argument(
        '--output',
        type=Path,
        help='Custom output path for the report'
    )

    args = parser.parse_args()

    try:
        # Archive mode
        if args.archive:
            archive_old_reports(days_to_keep=args.days, dry_run=args.dry_run)
            return 0

        # Generate mode (requires input file)
        if not args.input_file:
            parser.print_help()
            print("\nERROR: input_file is required (unless using --archive)", file=sys.stderr)
            return 1

        # Generate report
        output_path = generate_eod_report(
            args.input_file,
            output_path=args.output,
            dry_run=args.dry_run
        )

        if output_path:
            print("\nNext step: Upload this file to your Google Drive EOD folder.")

        return 0

    except FileNotFoundError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1
    except PermissionError as e:
        print(f"ERROR: Permission denied. Is the output file open in Word?", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"ERROR: Unexpected error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
