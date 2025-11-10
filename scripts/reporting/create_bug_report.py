#!/usr/bin/env python3
"""Create a DOCX bug report for Jira."""
import sys
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from common.docx_utils import get_report_output_path

def create_bug_report():
    """Create the bug report DOCX file."""
    doc = Document()

    # Title
    title = doc.add_heading('JIRA BUG REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bug emoji and subtitle
    subtitle = doc.add_heading('🐞 Bug Report Template – Jira', 1)

    # Title Section
    doc.add_heading('Title:', 2)
    p = doc.add_paragraph()
    p.add_run('Inconsistent wrong answer handling in "What\'s in the Classroom?" game - First question awards stars despite incorrect answer').bold = True

    # Priority
    doc.add_heading('Priority:', 2)
    priority = doc.add_paragraph()
    priority_run = priority.add_run('High')
    priority_run.bold = True
    priority_run.font.color.rgb = RGBColor(255, 0, 0)

    # Environment
    doc.add_heading('Environment:', 2)
    env_items = [
        'Platform: Hello Britannica Mobile App (iOS)',
        'App Version: 4.2.7',
        'Environment: Stage',
        'Testing Device: iPhone 12 Pro Max',
        'User Role: Student Account',
        f'Date Identified: November 10, 2025'
    ]
    for item in env_items:
        doc.add_paragraph(item, style='List Bullet')

    # Description
    doc.add_heading('Description:', 2)
    doc.add_paragraph(
        'The "What\'s in the Classroom?" activity under Level 1 > School & Education > Classroom Routines '
        'exhibits inconsistent behavior when students submit incorrect answers. The first question incorrectly '
        'awards 3 stars and allows progression despite displaying an error message for wrong answers. The second '
        'question behaves correctly by not awarding stars and blocking progression. This inconsistency undermines '
        'the educational integrity of the activity and creates confusion about answer correctness.'
    )

    # Steps to Reproduce
    doc.add_heading('Steps to Reproduce:', 2)

    steps = [
        '[Step 1 – Launch Hello Britannica Mobile App on iOS device and log in using student account credentials]',
        '',
        '[Step 2 – Navigate to: Level 1 > School & Education > Classroom Routines > Select "What\'s in the Classroom?" activity]',
        '',
        '[Step 3 – Observe the classroom image displayed for the first question]',
        '',
        '[Step 4 – Enter an intentionally incorrect answer in English describing the classroom (e.g., write "dog" when no dog is visible in the image)]',
        '',
        '[Step 5 – Submit the incorrect answer]',
        '',
        '[Step 6 – Observe that the game displays an error message indicating the answer is wrong, BUT incorrectly awards 3 stars and allows progression to the next question]',
        '',
        '[Step 7 – Proceed to the second question with the same classroom image but different prompt]',
        '',
        '[Step 8 – Enter an intentionally incorrect answer for the second question and submit]',
        '',
        '[Step 9 – Observe that the game correctly displays an error message, does NOT award stars, and does NOT allow progression]'
    ]
    for step in steps:
        if step:  # Only add non-empty lines
            p = doc.add_paragraph(step)
            if 'Step 6' in step:
                # Highlight the problematic step
                for run in p.runs:
                    if 'BUT incorrectly' in run.text or '3 stars' in run.text:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.bold = True
        else:
            doc.add_paragraph()  # Add blank line

    # Expected Behavior
    doc.add_heading('Expected Behavior:', 2)
    doc.add_paragraph('When a student submits an incorrect answer for ANY question in the activity:')
    expected = [
        'The game should display an error message indicating the answer is incorrect',
        'The game should NOT award any stars (0 stars)',
        'The game should NOT allow progression to the next question',
        'The student should be required to retry or be given appropriate feedback',
        'This behavior should be consistent across all questions in the activity'
    ]
    for item in expected:
        doc.add_paragraph(item, style='List Bullet')

    # Actual Behavior
    doc.add_heading('Actual Behavior:', 2)

    doc.add_heading('First Question:', 3)
    actual_first = [
        'Game displays error message: "Wrong answer" [CORRECT]',
        'Game awards 3 stars [INCORRECT - should award 0 stars]',
        'Game allows progression to next question [INCORRECT - should block progression]'
    ]
    for item in actual_first:
        p = doc.add_paragraph(item, style='List Bullet')
        if '[INCORRECT' in item:
            for run in p.runs:
                if '[INCORRECT' in run.text:
                    run.font.color.rgb = RGBColor(255, 0, 0)
        elif '[CORRECT]' in item:
            for run in p.runs:
                if '[CORRECT]' in run.text:
                    run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_heading('Second Question:', 3)
    actual_second = [
        'Game displays error message: "Wrong answer" [CORRECT]',
        'Game awards 0 stars [CORRECT]',
        'Game blocks progression [CORRECT]'
    ]
    for item in actual_second:
        p = doc.add_paragraph(item, style='List Bullet')
        if '[CORRECT]' in item:
            for run in p.runs:
                if '[CORRECT]' in run.text:
                    run.font.color.rgb = RGBColor(0, 128, 0)

    # Impact
    doc.add_heading('Impact:', 2)
    impacts = [
        'Educational Integrity: Students receive full credit (3 stars) for demonstrably incorrect answers on the first question, undermining the learning assessment',
        'User Confusion: Inconsistent feedback patterns may confuse students about whether their answers are actually correct or incorrect',
        'Progress Tracking: Parent/teacher dashboards may show inflated performance metrics due to unearned stars',
        'Learning Outcomes: Students may develop incorrect understanding of the material if rewarded for wrong answers',
        'System Consistency: Different validation logic between questions suggests a code defect that may affect other activities'
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style='List Bullet')

    # Suggested Fix
    doc.add_heading('Suggested Fix:', 2)
    fixes = [
        'Review the answer validation logic for the first question in the "What\'s in the Classroom?" activity',
        'Ensure the star-awarding function is properly gated by the answer validation result',
        'Verify that progression control is consistently applied across all questions',
        'Implement unit tests to ensure answer validation, star awarding, and progression control are properly synchronized',
        'Conduct regression testing on similar writing-based activities to identify if this affects other content'
    ]
    for i, fix in enumerate(fixes, 1):
        doc.add_paragraph(f'{i}. {fix}', style='List Number')

    # Notes
    doc.add_heading('Notes:', 2)
    notes = [
        'Video Evidence: Full reproduction captured on video, stored at: C:\\Users\\nico-\\AppData\\Roaming\\Microsoft\\Windows\\Libraries\\Videos.library-ms\\Mobile-writingquestion.mp4',
        'Reproducibility: 100% reproducible on every attempt',
        'Regression Risk: This may be a recent regression - check if recent updates modified answer validation logic',
        'Related Components: May affect other writing-based activities in Level 1 or across other levels',
        'Testing Recommendation: Test all writing prompt activities in the "Classroom Routines" section and similar game types',
        'User Reports: Verify if this has been reported by students/teachers in production'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    # Page break
    doc.add_page_break()

    # Jira Fields
    doc.add_heading('Bugs for Melingo / Options to Select When Creating Bug in Jira', 1)

    jira_fields = {
        'Project': 'Britannica ELL App (ELL)',
        'Work Type': 'Bug',
        'Priority': 'High (Major feature malfunction affecting assessment integrity and user experience)',
        'Components': 'Britannica',
        'Labels': 'ClasscodeApp',
        'Parent': 'Britannica QA - Q2',
        'Original Estimate': '2W'
    }

    for field, value in jira_fields.items():
        p = doc.add_paragraph()
        p.add_run(f'{field}: ').bold = True
        p.add_run(value)

    # Save the document
    output_path = Path(get_report_output_path('BUG_2025-11-10_Classroom_Game_Stars.docx'))
    doc.save(str(output_path))
    print(f"\n[OK] Bug report created successfully: {output_path}")
    return output_path

def main():
    try:
        output_path = create_bug_report()
        return 0
    except Exception as e:
        print(f"Error creating bug report: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
