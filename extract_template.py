#!/usr/bin/env python3
"""Extract Bug Template from DOCX file."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
except ImportError:
    print("python-docx not available, trying windows python")
    sys.exit(1)

def extract_template():
    doc_path = Path(__file__).parent / "documentation/templates/Bug Template for Jira.docx"
    doc = Document(doc_path)

    print("=== Bug Template for Jira ===\n")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Print with explicit encoding
            try:
                print(paragraph.text)
            except UnicodeEncodeError:
                print(paragraph.text.encode('utf-8', errors='replace').decode('utf-8'))

    if doc.tables:
        print("\n=== Tables ===")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                try:
                    print(row_text)
                except UnicodeEncodeError:
                    print(row_text.encode('utf-8', errors='replace').decode('utf-8'))

if __name__ == "__main__":
    extract_template()
